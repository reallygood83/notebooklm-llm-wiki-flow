"""Markdown to DOCX High-Precision Converter.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import List, Tuple

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

from .common import get_logger, step

LOG = get_logger("ax.docx")

HEADING_RE = re.compile(r"^(#{1,9})\s+(.*)$")
TABLE_SEP_RE = re.compile(r"^\|[\s:\-|]*\|$")
BOLD_SPLIT_RE = re.compile(r"(\*\*.*?\*\*)")

def _parse_md_table(lines: List[str], start_idx: int) -> Tuple[List[List[str]], int]:
    rows: List[List[str]] = []
    i = start_idx
    while i < len(lines):
        line = lines[i].strip()
        if not (line.startswith("|") and line.endswith("|")):
            break
        if TABLE_SEP_RE.match(line):
            i += 1
            continue
        cells = [c.strip() for c in line.split("|")[1:-1]]
        rows.append(cells)
        i += 1
    return rows, i

def _add_runs_with_bold(paragraph: Paragraph, text: str) -> None:
    for part in BOLD_SPLIT_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

def convert_md_to_docx(md_path: Path, docx_path: Path | None = None) -> Path:
    """Converts a Markdown file to a DOCX file with AX standards."""
    if not md_path.is_file():
        raise FileNotFoundError(f"Markdown file not found: {md_path}")

    out_path = docx_path or md_path.with_suffix(".docx")

    with step(LOG, "init-document"):
        doc = Document()
        normal = doc.styles["Normal"]
        # Standard administrative font for HWP/MS Word compatibility
        normal.font.name = "Malgun Gothic"
        normal.font.size = Pt(10)
        # font.name은 w:ascii/w:hAnsi만 설정한다. CJK 글자는 w:eastAsia 슬롯을
        # 따로 보지 않으면 Asian Theme Font(예: Batang)로 폴백된다. AX 산출물의
        # 주 콘텐츠가 한국어이므로 명시적으로 동일 폰트를 박아 둔다.
        rfonts = normal.element.get_or_add_rPr().get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), "Malgun Gothic")

    with step(LOG, "read-markdown"):
        lines = md_path.read_text(encoding="utf-8").splitlines()

    with step(LOG, "build-body"):
        i = 0
        while i < len(lines):
            line = lines[i].strip()

            if not line:
                doc.add_paragraph()
                i += 1
                continue

            heading_match = HEADING_RE.match(line)
            if heading_match:
                level = min(len(heading_match.group(1)), 9)
                text = heading_match.group(2)
                doc.add_heading(text, level=level)
                i += 1
                continue

            if line.startswith("|"):
                rows, next_i = _parse_md_table(lines, i)
                if rows:
                    cols = max(len(r) for r in rows)
                    rows = [r + [""] * (cols - len(r)) for r in rows]
                    table = doc.add_table(rows=len(rows), cols=cols)
                    table.style = "Table Grid"
                    for ri, row in enumerate(rows):
                        for ci, cell_text in enumerate(row):
                            clean = cell_text.replace("<br>", "\n").replace("**", "")
                            table.cell(ri, ci).text = clean
                    i = next_i
                    continue
                # `|`로 시작하지만 유효한 표 행이 아닌 라인(잘못된 표 row, |---| 단독 등).
                # 조용히 사라지지 않도록 텍스트 단락으로 보존하고 경고를 남긴다.
                LOG.warning(
                    "Pipe-prefixed line not parseable as table; rendering as text: %r",
                    line,
                )
                p = doc.add_paragraph()
                _add_runs_with_bold(p, line)
                i += 1
                continue

            p = doc.add_paragraph()
            _add_runs_with_bold(p, line)
            i += 1

    with step(LOG, "save-docx"):
        doc.save(str(out_path))

    return out_path
